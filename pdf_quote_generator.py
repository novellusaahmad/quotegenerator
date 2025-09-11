"""PDF, DOCX and Excel quote generation utilities.

Historically this module required the `reportlab` package at import time in
order to support PDF generation.  The DOCX helpers however do not rely on
`reportlab`, which meant simply importing this module would raise a
``ModuleNotFoundError`` when the dependency was missing.  As a consequence any
route that attempted to generate a Word document – such as the export feature
from ``calculator.html`` – resulted in a 500 error before the request even
reached the handler.

To make the DOCX generation more robust we now import ``reportlab`` lazily and
record whether it is available.  PDF helpers check this flag and return ``None``
when the library is absent, allowing callers to respond with a helpful error
message instead of causing an internal server error.  DOCX generation remains
unaffected and can function without the PDF dependency.
"""

import logging
import os
import re
import tempfile
from datetime import datetime
from decimal import Decimal

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ModuleNotFoundError:  # pragma: no cover - optional dependency
    REPORTLAB_AVAILABLE = False


BROTHER_FONT = "Brother 1816 Light"
BROTHER_STYLES = [
    "Normal",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Title",
    "List Bullet",
    "List Number",
]


logger = logging.getLogger(__name__)


def _apply_brother_font(doc):
    """Apply the Brother font and tightened spacing to common Word styles."""
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING

    for style_name in BROTHER_STYLES:
        try:
            style = doc.styles[style_name]
        except KeyError:
            # Skip any styles that do not exist in this document
            continue

        # Apply Brother font across different script categories
        font = style.font
        font.name = BROTHER_FONT
        r_fonts = style._element.rPr.rFonts
        r_fonts.set(qn("w:eastAsia"), BROTHER_FONT)
        r_fonts.set(qn("w:cs"), BROTHER_FONT)

        # Tighten line spacing to match design spec
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)

def generate_quote_pdf(quote_data, application_data=None):
    """Generate PDF quote document.

    Returns ``None`` if the optional ``reportlab`` dependency is not available.
    """
    if not REPORTLAB_AVAILABLE:
        return None

    # Create temporary file
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
        tmp_path = tmp_file.name

    # Create PDF document
    doc = SimpleDocTemplate(tmp_path, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    story.append(Paragraph("Novellus Finance - Loan Quote", title_style))
    story.append(Spacer(1, 12))
    
    # Quote details table
    quote_details = [
        ['Quote ID:', str(quote_data.get('id', 'N/A'))],
        ['Date:', datetime.now().strftime('%d/%m/%Y')],
        ['Gross Amount:', f"£{quote_data.get('gross_amount', 0):,.2f}"],
        ['Net Advance:', f"£{quote_data.get('net_advance', 0):,.2f}"],
        ['Interest Rate:', f"{quote_data.get('interest_rate', 0):.2f}%"],
        ['Loan Term:', f"{quote_data.get('loan_term', 0)} months"],
        ['Monthly Payment:', f"£{quote_data.get('monthly_payment', 0):,.2f}"],
        ['Total Interest:', f"£{quote_data.get('total_interest', 0):,.2f}"],
    ]
    
    table = Table(quote_details, colWidths=[2*inch, 3*inch])
    table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    story.append(table)
    story.append(Spacer(1, 24))
    
    # Footer
    footer_text = "This quote is subject to credit approval and full underwriting. Terms and conditions apply."
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        alignment=TA_CENTER
    )
    story.append(Paragraph(footer_text, footer_style))
    
    # Build PDF
    doc.build(story)
    
    # Read the PDF content
    with open(tmp_path, 'rb') as f:
        pdf_content = f.read()
    
    # Clean up temporary file
    os.unlink(tmp_path)
    
    return pdf_content

def generate_professional_quote_docx(quote_data, application_data=None):
    """Generate professional DOCX quote document"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ModuleNotFoundError:
        # Required dependency is missing. Return ``None`` so callers can
        # provide a helpful error message rather than raising an
        # unhandled exception which results in a generic 500 response.
        return None

    doc = Document()
    _apply_brother_font(doc)

    # Determine currency-specific logo
    currency = quote_data.get('currency', 'GBP')
    logo_map = {
        'GBP': 'novellus_logo_gbp.svg',
        'EUR': 'novellus_logo_eur.svg',
    }
    logo_filename = logo_map.get(currency, 'novellus_logo_gbp.svg')
    logo_path = os.path.join(os.path.dirname(__file__), 'static', logo_filename)

    # Add header with logo
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(logo_path):
        header_para.add_run().add_picture(logo_path, width=Inches(1.3))

    # Add title
    title = doc.add_heading('Novellus Finance - Loan Quote', 0)
    title.alignment = 1  # Center alignment

    # Add quote details
    doc.add_heading('Quote Details', level=1)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    details = [
        ('Quote ID:', str(quote_data.get('id', 'N/A'))),
        ('Date:', datetime.now().strftime('%d/%m/%Y')),
        ('Gross Amount:', f"£{quote_data.get('gross_amount', 0):,.2f}"),
        ('Net Advance:', f"£{quote_data.get('net_advance', 0):,.2f}"),
        ('Interest Rate:', f"{quote_data.get('interest_rate', 0):.2f}%"),
        ('Loan Term:', f"{quote_data.get('loan_term', 0)} months"),
        ('Monthly Payment:', f"£{quote_data.get('monthly_payment', 0):,.2f}"),
        ('Total Interest:', f"£{quote_data.get('total_interest', 0):,.2f}"),
    ]
    
    for i, (key, value) in enumerate(details):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = value

    # Footer with centered logo
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ''
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_path):
        footer_para.add_run().add_picture(logo_path, width=Inches(1))

    for run in footer_para.runs:
        run.font.size = Pt(8)

    # Save to temporary file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        doc.save(tmp_file.name)
        tmp_path = tmp_file.name
    
    # Read the content
    with open(tmp_path, 'rb') as f:
        docx_content = f.read()
    
    # Clean up
    os.unlink(tmp_path)
    
    return docx_content


def generate_loan_summary_docx(loan, extra_fields=None):
    """Generate DOCX loan summary report.

    Parameters
    ----------
    loan: LoanSummary
        Database loan summary instance containing core data.
    extra_fields: dict, optional
        Additional user supplied fields gathered from the modal form.  The
        dictionary may include a list of ``note_templates`` (or legacy
        ``notes``).  Each note may specify a ``placeholder_map`` used to map
        tokens in the note text to keys within a unified context built from the
        loan and the extra fields.  Any tokens without corresponding values are
        replaced with an empty string and a warning is logged.  Legacy behaviour
        where tokens were substituted directly from ``extra_fields`` keys is
        deprecated.
    """
    extra_fields = extra_fields or {}
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
    except ModuleNotFoundError:
        return None
    import tempfile

    doc = Document()
    _apply_brother_font(doc)

    # Determine currency-specific logo
    currency = getattr(loan, 'currency', 'GBP')
    logo_map = {

        'GBP': 'novellus_logo_gbp.png',
        'EUR': 'novellus_logo_eur.png',
    }
    logo_filename = logo_map.get(currency, 'novellus_logo_gbp.png')

    logo_path = os.path.join(os.path.dirname(__file__), 'static', logo_filename)

    def _add_hyperlink(paragraph, url, text):
        """Add a hyperlink to a docx paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")  # 8pt font size
        rPr.append(sz)
        new_run.append(rPr)
        w_text = OxmlElement("w:t")
        w_text.text = text
        new_run.append(w_text)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    # Header with logo
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(logo_path):
        try:
            header_para.add_run().add_picture(logo_path, width=Inches(1.3))
        except Exception as exc:  # pragma: no cover - best effort logging
            logging.getLogger(__name__).warning(
                "Unable to load header logo %s: %s", logo_path, exc
            )

    client_name = extra_fields.get('client_name', '[•]')
    doc.add_paragraph(f"Dear {client_name},")
    subjects = []
    if extra_fields.get('include_valuation', True):
        subjects.append('valuation')
    if extra_fields.get('include_planning_appraisal', True):
        subjects.append('planning appraisal')
    if extra_fields.get('include_qs_appraisal', True):
        subjects.append('QS appraisal')
    if extra_fields.get('include_due_diligence', True):
        subjects.append('due diligence')
    if extra_fields.get('include_legals', True):
        subjects.append('legals')

    def _roman(n):
        numerals = ['i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x']
        return numerals[n - 1]

    if subjects:
        parts = [f"({_roman(i + 1)}) {text}" for i, text in enumerate(subjects)]
        if len(parts) > 1:
            clause = ', '.join(parts[:-1]) + ' and ' + parts[-1]
        else:
            clause = parts[0]
        doc.add_paragraph(
            f"Further to our correspondence, please see below our high-level terms subject to {clause}:"
        )
    else:
        doc.add_paragraph(
            "Further to our correspondence, please see below our high-level terms:"
        )

    # Helper to add bullet paragraphs with optional bold segments
    def _add_numbered(parts):
        if not any(t.strip() for t, _ in parts):
            return
        p = doc.add_paragraph(style='List Number')
        for text, bold in parts:
            run = p.add_run(text)
            run.bold = bold
            run.font.color.rgb = RGBColor(0, 0, 0)

    property_address = extra_fields.get('property_address')
    if property_address:
        addresses = [
            re.sub(r'^\d+\.\s*', '', a.strip())
            for a in re.split(r';|\n', property_address)
            if a.strip()
        ]
        if len(addresses) > 1:
            doc.add_paragraph("Property Addresses:")
            for idx, addr in enumerate(addresses, 1):
                p = doc.add_paragraph()
                num_run = p.add_run(f"{idx}. ")
                num_run.font.color.rgb = RGBColor(0, 0, 0)
                run = p.add_run(addr)
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif addresses:
            doc.add_paragraph(f"Property Address: {addresses[0]}")

    # Table with header row whose color depends on currency
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    col_widths = (Inches(3.5), Inches(1.2), Inches(2.3))
    for width, column in zip(col_widths, table.columns):
        column.width = width
        for cell in column.cells:
            cell.width = width

    # Remove inner borders, keeping only the outer grid
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.append(tbl_pr)
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    # Ensure outer borders exist
    for edge in ("top", "left", "bottom", "right"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
    # Remove inside borders
    for edge in ("insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "none")

    currency_symbol = '€' if getattr(loan, 'currency', 'GBP') == 'EUR' else '£'
    arr_fee_pct = f"{float(getattr(loan, 'arrangement_fee_percentage', 0) or 0):.2f}%"
    ltv_ratio = getattr(loan, 'ltv_ratio', getattr(loan, 'start_ltv', 0))

    term_val = getattr(loan, 'loan_term', 0) or 0
    term_text = str(term_val)
    interest_term_text = f"{term_val} Month{'s' if term_val != 1 else ''}"

    rows = [
        ("Valuation", "", f"{currency_symbol}{float(getattr(loan, 'property_value', 0) or 0):,.2f}"),
        ("Gross Amount", "", f"{currency_symbol}{float(getattr(loan, 'gross_amount', 0) or 0):,.2f}"),
        ("Term (Months)", term_text, ""),
        (
            "Arrangement Fee",
            arr_fee_pct,
            f"{currency_symbol}{float(getattr(loan, 'arrangement_fee', 0) or 0):,.2f}",
        ),
        (
            "Legal Costs & Title Insurance*",
            "",
            f"{currency_symbol}{float((getattr(loan, 'legal_costs', 0) or 0) + (getattr(loan, 'title_insurance', 0) or 0)):,.2f}",
        ),
        (
            "Number Months (Interest)",
            interest_term_text,
            f"{currency_symbol}{float(getattr(loan, 'total_interest', 0) or 0):,.2f}",
        ),
        (
            "Day 1 Net Advance",
            "",
            f"{currency_symbol}{float((getattr(loan, 'day_1_advance', None) or getattr(loan, 'net_advance', 0) or 0)):,.2f}",
        ),
        (
            "Total Net Advance",
            "",
            f"{currency_symbol}{float(getattr(loan, 'total_net_advance', 0) or 0):,.2f}",
        ),
    ]

    def _is_numeric(text):
        """Return True if ``text`` represents a number, currency or percentage."""
        if text is None:
            return False
        cleaned = str(text).strip()
        if cleaned == "":
            return False
        cleaned = cleaned.replace(",", "").replace("£", "").replace("€", "").replace("%", "")
        try:
            float(cleaned)
            return True
        except ValueError:
            match = re.match(r"^[+-]?\d+(?:\.\d+)?", cleaned)
            return bool(match)

    # Create heading row
    heading_cell = table.cell(0, 0)
    heading_cell.text = "Loan Summary"
    heading_cell.merge(table.cell(0, 2))
    heading_para = heading_cell.paragraphs[0]
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading_para.runs[0]
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 0)

    # Apply currency-specific color to header
    header_color = "509664" if currency == 'EUR' else "#D1BE5D"
    tc_pr = heading_cell._tc.get_or_add_tcPr()
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), header_color))
    tc_pr.append(shd)

    # Populate remaining rows with alignment based on content type
    for i, (c1, c2, c3) in enumerate(rows, start=1):
        table.cell(i, 0).text = c1
        p0 = table.cell(i, 0).paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT

        table.cell(i, 1).text = c2
        p1 = table.cell(i, 1).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT if _is_numeric(c2) else WD_ALIGN_PARAGRAPH.LEFT

        table.cell(i, 2).text = c3
        p2 = table.cell(i, 2).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT if _is_numeric(c3) else WD_ALIGN_PARAGRAPH.LEFT

        for col in (1, 2):
            for run in table.cell(i, col).paragraphs[0].runs:
                run.font.bold = True

    # Shade alternating rows
    for row_idx, row in enumerate(table.rows[1:], start=1):
        shade = "F8F9FA" if row_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shade))
            tc_pr.append(shd)

    # Build a unified context from the loan attributes and any additional
    # ``extra_fields``.  Values are stored using lower-case keys to allow
    # case-insensitive lookups.  Nested dictionaries or objects are stored
    # directly so they can be traversed via ``_resolve_path`` when
    # substituting tokens.
    def _flatten(value):
        return float(value) if isinstance(value, Decimal) else value

    def _resolve_path(obj, path):
        """Resolve a dotted ``path`` against ``obj``.

        Each segment in the path is looked up case-insensitively on either
        dictionaries or object attributes.  ``None`` is returned if any segment
        is missing."""

        current = obj
        for segment in path.split('.'):
            if isinstance(current, dict):
                if segment in current:
                    current = current[segment]
                else:
                    current = current.get(segment.lower())
            else:
                current = getattr(current, segment, getattr(current, segment.lower(), None))
            if current is None:
                return None
        return _flatten(current)

    # Ensure the formatted snapshot is loaded so fields can be referenced via
    # ``loan_data`` in placeholder mappings.
    loan_data_obj = getattr(loan, 'loan_data', None)

    context = {}
    for attr, value in vars(loan).items():
        if attr.startswith('_'):
            continue
        context[attr.lower()] = _flatten(value)
    if loan_data_obj is not None:
        context['loan_data'] = loan_data_obj
    for key, value in extra_fields.items():
        if isinstance(value, list):
            continue
        context[key.lower()] = _flatten(value)

    max_ltv = extra_fields.get('max_ltv') or float(ltv_ratio or 0)
    context.setdefault('max_ltv', max_ltv)
    context.setdefault('ltv', max_ltv)

    token_pattern = re.compile(r"\[([^\]]+)\]", re.IGNORECASE)

    def _replace_tokens(text, placeholder_map=None):
        cleaned_map = {}
        for key, value in (placeholder_map or {}).items():
            norm_key = str(key).strip().strip("[]").lower()
            if isinstance(value, str):
                norm_val = value.strip()
                if norm_val.lower().startswith("loan_summary."):
                    norm_val = norm_val.split(".", 1)[1]
                norm_val = norm_val.lower()
            else:
                norm_val = value
            cleaned_map[norm_key] = norm_val

        def repl(match):
            token_name = match.group(1)
            ctx_key = cleaned_map.get(token_name.lower(), token_name)
            if isinstance(ctx_key, str) and ctx_key.lower().startswith("loan_summary."):
                ctx_key = ctx_key.split(".", 1)[1]
            value = _resolve_path(context, ctx_key)
            if value in (None, ""):
                logger.warning("Missing value for placeholder %s", token_name)
                return ""
            return str(value)

        return token_pattern.sub(repl, text)


    note_templates = (
        extra_fields.get('note_templates')
        or extra_fields.get('notes')
        or []
    )

    security_notes, salient_notes, other_notes = [], [], []
    for nt in note_templates:
        group = ''
        if isinstance(nt, dict):
            group = (nt.get('group') or '').lower()
        if group == 'security':
            security_notes.append(nt)
        elif group in ('salient point', 'salient points'):
            salient_notes.append(nt)
        else:
            other_notes.append(nt)

    sections = []
    if security_notes:
        sections.append(('Security', security_notes))
    if salient_notes:
        sections.append(('Salient Point', salient_notes))
    if other_notes:
        sections.append(('Conditions', other_notes))

    for heading, bullets in sections:
        hp = doc.add_heading(_replace_tokens(heading), level=1)
        for run in hp.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        for idx, bullet in enumerate(bullets, 1):
            placeholder_map = {}
            text = bullet
            if isinstance(bullet, dict):
                placeholder_map = bullet.get('placeholder_map') or {}
                text = bullet.get('text', '')
            if isinstance(text, list):
                parts = [(_replace_tokens(t, placeholder_map), b) for t, b in text]
            else:
                parts = [(_replace_tokens(text, placeholder_map), False)]
            if not any(t.strip() for t, _ in parts):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            num_run = p.add_run(f"{idx}. ")
            num_run.font.color.rgb = RGBColor(0, 0, 0)
            for t, b in parts:
                run = p.add_run(t)
                run.bold = b
                run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("Yours sincerely, [or faithfully if Dear Sir],")
    doc.add_paragraph("[•]")
    doc.add_paragraph("For and on behalf of")
    doc.add_paragraph("Novellus Finance Limited")

    # Footer with center aligned text and email link
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ''
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


    if currency == 'EUR':
        footer_lines = [
            "Novellus Finance Limited trading as Novellus Finance is registered in Ireland. Company Reg. No 710946.",
            "100 St Stephen’s Green, Dublin, D02 EP84 | +353 1531 4237 | info@novellusfinance.com |",
            "Novellus Finance Limited is not regulated by the Central Bank of Ireland.",
        ]
    else:
        footer_lines = [
            "T Bromley, 15 London Road, Bromley, Kent BR1 1DE | 0203 397 4871  |  info@novellusfinance.com  | novellusfinance.com",
            "Novellus Limited trading as Novellus Finance is registered in England & Wales Company Reg. No 10790634",
            "Novellus Limited is not regulated by the Financial Conduct Authority",
        ]

    for i, line in enumerate(footer_lines):
        if i:
            footer_para.add_run().add_break()
        if "info@novellusfinance.com" in line:
            before, after = line.split("info@novellusfinance.com")
            footer_para.add_run(before)
            _add_hyperlink(
                footer_para,
                "mailto:info@novellusfinance.com",
                "info@novellusfinance.com",
            )
            footer_para.add_run(after)
        else:
            footer_para.add_run(line)

    for run in footer_para.runs:
        run.font.size = Pt(8)
        if currency == 'EUR':
            run.font.color.rgb = RGBColor(0x50, 0x96, 0x64)
        else:
            run.font.color.rgb = RGBColor(0, 0, 0)


    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        doc.save(tmp_file.name)
        tmp_path = tmp_file.name

    with open(tmp_path, 'rb') as f:
        docx_content = f.read()

    os.unlink(tmp_path)
    return docx_content
