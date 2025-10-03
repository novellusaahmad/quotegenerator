import io
import os
from docx import Document

os.environ["DATABASE_URL"] = "sqlite:///test.db"

from app import app, db
from models import LoanSummary, ReportFields, LoanNote


def _setup_loan_and_note(placeholder_map=None):
    """Reset DB and create a loan with a note using MAX_LTV placeholder."""
    with app.app_context():
        db.drop_all()
        db.create_all()
        loan = LoanSummary(loan_name="TestLoan", loan_type="bridge")
        db.session.add(loan)
        note = LoanNote(
            group="General",
            name="Max LTV is [MAX_LTV]%",
            placeholder_map=placeholder_map
            or {"MAX_LTV": "report_fields.max_ltv"},
            add_flag=True,
        )
        db.session.add(note)
        db.session.commit()
        return loan.id


def _extract_text(doc_bytes):
    doc = Document(io.BytesIO(doc_bytes))
    text = []
    for p in doc.paragraphs:
        text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return "\n".join(text)


def test_download_summary_docx_post_uses_report_fields():
    loan_id = _setup_loan_and_note()
    client = app.test_client()
    payload = {"max_ltv": 70}
    res = client.post(f"/loan/{loan_id}/summary-docx", json=payload)
    assert res.status_code == 200
    text = _extract_text(res.data)
    assert "Max LTV is 70%" in text
    assert "[MAX_LTV]" not in text


def test_download_summary_docx_get_uses_saved_report_fields():
    loan_id = _setup_loan_and_note()
    with app.app_context():
        rf = ReportFields(loan_id=loan_id, max_ltv=55)
        db.session.add(rf)
        db.session.commit()
        assert LoanNote.query.count() == 1
    client = app.test_client()
    res = client.get(f"/loan/{loan_id}/summary-docx")
    assert res.status_code == 200
    text = _extract_text(res.data)
    assert "Max LTV is 55.0%" in text
    assert "[MAX_LTV]" not in text


def test_report_fields_post_updates_snapshot():
    """Report fields save should refresh LoanData snapshot for placeholders."""
    loan_id = _setup_loan_and_note({"MAX_LTV": "max_ltv"})
    with app.app_context():
        note_id = LoanNote.query.first().id
    client = app.test_client()
    payload = {"max_ltv": 72, "note_ids": [note_id]}
    res = client.post(f"/loan/{loan_id}/report-fields", json=payload)
    assert res.status_code == 200
    res = client.get(f"/loan/{loan_id}/summary-docx")
    assert res.status_code == 200
    text = _extract_text(res.data)
    assert "Max LTV is 72.0%" in text
    assert "[MAX_LTV]" not in text


def test_loan_notes_grouped_headings_and_numbering():
    with app.app_context():
        db.drop_all()
        db.create_all()
        loan = LoanSummary(loan_name="TestLoan", loan_type="bridge")
        db.session.add(loan)
        notes = [
            LoanNote(group="Security", name="Security note", add_flag=True),
            LoanNote(group="Salient Point", name="Salient note", add_flag=True),
            LoanNote(group="General", name="General note", add_flag=True),
        ]
        db.session.add_all(notes)
        db.session.commit()
        loan_id = loan.id
    client = app.test_client()
    res = client.get(f"/loan/{loan_id}/summary-docx")
    assert res.status_code == 200
    text_lines = _extract_text(res.data).splitlines()
    assert "Security" in text_lines
    assert "Salient Point" in text_lines
    assert "Conditions" in text_lines
    sec_idx = text_lines.index("Security")
    assert text_lines[sec_idx + 1].startswith("1. ")
    sal_idx = text_lines.index("Salient Point")
    assert text_lines[sal_idx + 1].startswith("1. ")
    cond_idx = text_lines.index("Conditions")
    assert text_lines[cond_idx + 1].startswith("1. ")


def test_docx_respects_note_sort_order():
    with app.app_context():
        db.drop_all()
        db.create_all()
        loan = LoanSummary(loan_name="TestLoan", loan_type="bridge")
        db.session.add(loan)
        notes = [
            LoanNote(
                group="General",
                name="Second general note",
                add_flag=True,
                sort_order=1,
            ),
            LoanNote(
                group="General",
                name="First general note",
                add_flag=True,
                sort_order=0,
            ),
            LoanNote(
                group="General",
                name="Third general note",
                add_flag=True,
                sort_order=2,
            ),
        ]
        db.session.add_all(notes)
        db.session.commit()
        loan_id = loan.id

    client = app.test_client()
    res = client.get(f"/loan/{loan_id}/summary-docx")
    assert res.status_code == 200
    text_lines = _extract_text(res.data).splitlines()
    bullet_lines = [
        line.strip()
        for line in text_lines
        if "general note" in line.lower()
    ]
    assert bullet_lines[:3] == [
        "1. First general note",
        "2. Second general note",
        "3. Third general note",
    ]


def test_reorder_endpoint_resequences_duplicate_sort_orders():
    with app.app_context():
        db.drop_all()
        db.create_all()
        notes = [
            LoanNote(group="General", name="First", add_flag=True, sort_order=0),
            LoanNote(group="General", name="Second", add_flag=True, sort_order=0),
            LoanNote(group="General", name="Third", add_flag=True, sort_order=0),
        ]
        db.session.add_all(notes)
        db.session.commit()
        note_ids = [note.id for note in notes]

    client = app.test_client()
    res = client.post(
        f"/loan-notes/{note_ids[0]}/reorder",
        json={"direction": "down"},
    )
    assert res.status_code == 200
    data = res.get_json()
    assert data["success"] is True
    assert [item["sort_order"] for item in data["order"]] == [0, 1, 2]

    with app.app_context():
        ordered = (
            LoanNote.query.filter_by(group="General", deleted_at=None)
            .order_by(LoanNote.sort_order, LoanNote.id)
            .all()
        )
        assert [n.sort_order for n in ordered] == [0, 1, 2]
        assert [n.id for n in ordered] == [note_ids[1], note_ids[0], note_ids[2]]


def test_multiple_property_addresses_numbered():
    with app.app_context():
        db.drop_all()
        db.create_all()
        loan = LoanSummary(loan_name="TestLoan", loan_type="bridge")
        db.session.add(loan)
        db.session.commit()
        loan_id = loan.id

    client = app.test_client()
    payload = {
        "property_address": "123 Example Street\n456 Another Ave",
    }
    res = client.post(f"/loan/{loan_id}/summary-docx", json=payload)
    assert res.status_code == 200
    text_lines = _extract_text(res.data).splitlines()
    assert "1. 123 Example Street" in text_lines
    assert "2. 456 Another Ave" in text_lines


def test_pre_numbered_property_addresses_not_double_numbered():
    with app.app_context():
        db.drop_all()
        db.create_all()
        loan = LoanSummary(loan_name="TestLoan", loan_type="bridge")
        db.session.add(loan)
        db.session.commit()
        loan_id = loan.id

    client = app.test_client()
    payload = {
        "property_address": "1. 123 Example Street\n2. 456 Another Ave",
    }
    res = client.post(f"/loan/{loan_id}/summary-docx", json=payload)
    assert res.status_code == 200
    text_lines = _extract_text(res.data).splitlines()
    assert "1. 123 Example Street" in text_lines
    assert "2. 456 Another Ave" in text_lines
    assert all(
        not line.startswith("1. 1.") and not line.startswith("2. 2.")
        for line in text_lines
    )
