import io
from types import SimpleNamespace

from docx import Document

from pdf_quote_generator import generate_loan_summary_docx


def extract_text(doc_bytes):
    document = Document(io.BytesIO(doc_bytes))
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return "\n".join(text)


def test_placeholder_map_replaces_token():
    loan = SimpleNamespace(
        currency='GBP',
        arrangement_fee_percentage=0,
        arrangement_fee=0,
        property_value=0,
        gross_amount=0,
        loan_term=0,
        legal_costs=0,
        title_insurance=0,
        total_interest=0,
        day_1_advance=0,
        net_advance=0,
        total_net_advance=0,
        ltv_ratio=0,
        start_ltv=0,
    )
    extra_fields = {
        'max_ltv': 70,
        'note_templates': [
            {
                'text': 'Max LTV is [MAX_LTV]%',
                'placeholder_map': {'MAX_LTV': 'max_ltv'},
            }
        ],
    }

    doc_bytes = generate_loan_summary_docx(loan, extra_fields)
    text = extract_text(doc_bytes)

    assert 'Max LTV is 70%' in text
    assert '[MAX_LTV]' not in text


def test_placeholder_map_nested_path():
    loan = SimpleNamespace(
        currency='GBP',
        arrangement_fee_percentage=0,
        arrangement_fee=0,
        property_value=0,
        gross_amount=0,
        loan_term=0,
        legal_costs=0,
        title_insurance=0,
        total_interest=0,
        day_1_advance=0,
        net_advance=0,
        total_net_advance=0,
        ltv_ratio=0,
        start_ltv=0,
    )

    extra_fields = {
        'report_fields': {'max_ltv': 60},
        'note_templates': [
            {
                'text': 'Max LTV is [MAX_LTV]%',
                'placeholder_map': {'MAX_LTV': 'report_fields.max_ltv'},
            }
        ],
    }

    doc_bytes = generate_loan_summary_docx(loan, extra_fields)
    text = extract_text(doc_bytes)

    assert 'Max LTV is 60%' in text
    assert '[MAX_LTV]' not in text


def test_placeholder_map_report_fields_block():
    loan = SimpleNamespace(
        currency='GBP',
        arrangement_fee_percentage=0,
        arrangement_fee=0,
        property_value=0,
        gross_amount=0,
        loan_term=0,
        legal_costs=0,
        title_insurance=0,
        total_interest=0,
        day_1_advance=0,
        net_advance=0,
        total_net_advance=0,
        ltv_ratio=0,
        start_ltv=0,
        report_fields=SimpleNamespace(max_ltv=55),
    )

    extra_fields = {
        'max_ltv': 55,
        'report_fields': {'max_ltv': 55},
        'note_templates': [
            {
                'text': 'Financial Covenants Maximum LTV [MAX_LTV]%',
                'placeholder_map': {'MAX_LTV': 'report_fields.max_ltv'},
            }
        ],
    }

    doc_bytes = generate_loan_summary_docx(loan, extra_fields)
    text = extract_text(doc_bytes)

    assert 'Financial Covenants Maximum LTV 55%' in text
    assert '[MAX_LTV]' not in text



def test_placeholder_map_bracket_keys_with_prefix():
    loan = SimpleNamespace(
        currency='GBP',
        arrangement_fee_percentage=5,
        arrangement_fee=1000,
        property_value=0,
        gross_amount=0,
        loan_term=0,
        legal_costs=0,
        title_insurance=0,
        total_interest=0,
        day_1_advance=0,
        net_advance=0,
        total_net_advance=0,
        ltv_ratio=0,
        start_ltv=0,
    )

    extra_fields = {
        'note_templates': [
            {
                'text': 'The arrangement fee is [ARRANGEMENT_FEE] i.e. [ARRANGEMENT_FEE_PER]% of the gross loan',
                'placeholder_map': {
                    '[ARRANGEMENT_FEE]': 'loan_summary.arrangement_fee',
                    '[ARRANGEMENT_FEE_PER]': 'loan_summary.arrangement_fee_percentage',
                },
            }
        ]
    }

    doc_bytes = generate_loan_summary_docx(loan, extra_fields)
    text = extract_text(doc_bytes)

    assert 'The arrangement fee is 1000 i.e. 5%' in text
    assert '[ARRANGEMENT_FEE]' not in text
    assert '[ARRANGEMENT_FEE_PER]' not in text


def test_table_alignment_and_borders():
    loan = SimpleNamespace(
        currency='GBP',
        arrangement_fee_percentage=5,
        arrangement_fee=1000,
        property_value=200000,
        gross_amount=150000,
        loan_term=13,
        legal_costs=3000,
        title_insurance=0,
        total_interest=12986.30,
        day_1_advance=82013.7,
        net_advance=82013.7,
        total_net_advance=82013.7,
        ltv_ratio=0,
        start_ltv=0,
    )

    doc_bytes = generate_loan_summary_docx(loan, {})
    document = Document(io.BytesIO(doc_bytes))
    table = document.tables[0]

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    # Numeric and percentage fields should be right aligned
    assert table.cell(3, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert table.cell(4, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert table.cell(1, 2).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    # Text fields left aligned
    assert table.cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert table.cell(6, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT

    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    from docx.oxml.ns import qn
    inside_h = tbl_borders.find(qn('w:insideH'))
    inside_v = tbl_borders.find(qn('w:insideV'))
    assert inside_h is not None and inside_h.get(qn('w:val')) == 'none'
    assert inside_v is not None and inside_v.get(qn('w:val')) == 'none'

